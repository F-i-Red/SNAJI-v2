"""
Processador de Documentos do SNAJI.

Extrai texto de ficheiros PDF e Word para análise jurídica.
Funciona de forma assíncrona para não bloquear o servidor.

Suportado:
- PDF (.pdf) — via PyPDF2
- Word (.docx) — via python-docx
- Texto simples (.txt)
- Tamanho máximo: 10MB por ficheiro
"""

from __future__ import annotations
import io
import re
from dataclasses import dataclass
from enum import Enum
from pathlib import Path

import structlog

logger = structlog.get_logger(__name__)

MAX_BYTES = 10 * 1024 * 1024  # 10MB


class TipoFicheiro(str, Enum):
    PDF  = "pdf"
    DOCX = "docx"
    TXT  = "txt"
    DESCONHECIDO = "desconhecido"


@dataclass
class DocumentoExtraido:
    nome_original: str
    tipo: TipoFicheiro
    texto: str
    num_paginas: int
    num_caracteres: int
    avisos: list[str]


def _detectar_tipo(nome: str, conteudo: bytes) -> TipoFicheiro:
    """Detecta o tipo pelo magic bytes e extensão — não confia apenas na extensão."""
    ext = Path(nome).suffix.lower()
    if conteudo[:4] == b'%PDF':
        return TipoFicheiro.PDF
    if conteudo[:2] == b'PK' and ext == '.docx':
        return TipoFicheiro.DOCX
    if ext in ('.txt', '.md'):
        return TipoFicheiro.TXT
    return TipoFicheiro.DESCONHECIDO


def _extrair_pdf(conteudo: bytes) -> tuple[str, int, list[str]]:
    """Extrai texto de PDF página a página."""
    try:
        import PyPDF2
        leitor = PyPDF2.PdfReader(io.BytesIO(conteudo))
        paginas = []
        avisos = []
        for i, pagina in enumerate(leitor.pages):
            try:
                texto = pagina.extract_text() or ""
                if texto.strip():
                    paginas.append(texto)
            except Exception:
                avisos.append(f"Página {i+1} não pôde ser extraída")
        texto_total = "\n\n".join(paginas)
        return texto_total, len(leitor.pages), avisos
    except ImportError:
        return "", 0, ["PyPDF2 não instalado — instalar com: pip install PyPDF2"]
    except Exception as e:
        return "", 0, [f"Erro ao processar PDF: {e}"]


def _extrair_docx(conteudo: bytes) -> tuple[str, int, list[str]]:
    """Extrai texto de ficheiro Word (.docx)."""
    try:
        import docx
        doc = docx.Document(io.BytesIO(conteudo))
        paragrafos = [p.text for p in doc.paragraphs if p.text.strip()]
        texto = "\n\n".join(paragrafos)
        return texto, len(doc.paragraphs), []
    except ImportError:
        return "", 0, ["python-docx não instalado — instalar com: pip install python-docx"]
    except Exception as e:
        return "", 0, [f"Erro ao processar DOCX: {e}"]


def _limpar_texto(texto: str) -> str:
    """Remove ruído comum em textos extraídos de PDF."""
    # Remove múltiplas linhas em branco
    texto = re.sub(r'\n{4,}', '\n\n\n', texto)
    # Remove espaços no início de linha
    texto = re.sub(r'^ +', '', texto, flags=re.MULTILINE)
    # Remove caracteres de controlo excepto newline e tab
    texto = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', texto)
    return texto.strip()


class ProcessadorDocumentos:
    """
    Extrai texto de documentos para análise jurídica.
    Stateless — cada chamada é independente.
    """

    def processar(self, nome: str, conteudo: bytes) -> DocumentoExtraido:
        """
        Processa um documento e devolve o texto extraído.
        Nunca lança excepção — devolve avisos em caso de erro.
        """
        log = logger.bind(nome=nome, bytes=len(conteudo))

        if len(conteudo) > MAX_BYTES:
            return DocumentoExtraido(
                nome_original=nome,
                tipo=TipoFicheiro.DESCONHECIDO,
                texto="",
                num_paginas=0,
                num_caracteres=0,
                avisos=[f"Ficheiro demasiado grande ({len(conteudo) // 1024}KB). Máximo: 10MB."],
            )

        tipo = _detectar_tipo(nome, conteudo)
        log.info("doc.processar", tipo=tipo.value)

        avisos: list[str] = []
        num_paginas = 1

        if tipo == TipoFicheiro.PDF:
            texto, num_paginas, avisos = _extrair_pdf(conteudo)
        elif tipo == TipoFicheiro.DOCX:
            texto, num_paginas, avisos = _extrair_docx(conteudo)
        elif tipo == TipoFicheiro.TXT:
            try:
                texto = conteudo.decode('utf-8', errors='replace')
            except Exception as e:
                texto = ""
                avisos.append(f"Erro de codificação: {e}")
        else:
            texto = ""
            avisos.append(f"Tipo de ficheiro não suportado: {Path(nome).suffix}")

        texto = _limpar_texto(texto)

        if not texto and not avisos:
            avisos.append("Não foi possível extrair texto. O ficheiro pode estar protegido ou ser uma imagem.")

        log.info("doc.processado", chars=len(texto), paginas=num_paginas, avisos=len(avisos))

        return DocumentoExtraido(
            nome_original=nome,
            tipo=tipo,
            texto=texto,
            num_paginas=num_paginas,
            num_caracteres=len(texto),
            avisos=avisos,
        )
